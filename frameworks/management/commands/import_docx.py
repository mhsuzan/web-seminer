"""
Management command to import Knowledge Graph quality frameworks data from a Word document.

Usage:
    python manage.py import_docx path/to/document.docx

The document should contain tables or structured data with framework information.
This is a basic implementation that can be extended based on the actual document structure.
"""
from django.core.management.base import BaseCommand, CommandError
from django.db import transaction
from docx import Document
import re
from frameworks.models import Framework, Criterion, Definition


class Command(BaseCommand):
    help = 'Import Knowledge Graph quality frameworks from a Word document'

    def add_arguments(self, parser):
        parser.add_argument('docx_file', type=str, help='Path to the Word document (.docx)')
        parser.add_argument(
            '--dry-run',
            action='store_true',
            help='Run without actually saving data to database',
        )

    def handle(self, *args, **options):
        docx_path = options['docx_file']
        dry_run = options['dry_run']

        try:
            doc = Document(docx_path)
            self.stdout.write(self.style.SUCCESS(f'Opened document: {docx_path}'))
            
            if dry_run:
                self.stdout.write(self.style.WARNING('DRY RUN MODE - No data will be saved'))
            
            # Parse the document
            frameworks_data = self.parse_document(doc)
            
            self.stdout.write(self.style.SUCCESS(f'Found {len(frameworks_data)} frameworks'))
            
            if not dry_run:
                with transaction.atomic():
                    imported_count = self.import_frameworks(frameworks_data)
                    self.stdout.write(
                        self.style.SUCCESS(f'Successfully imported {imported_count} frameworks')
                    )
            else:
                self.stdout.write('Would import:')
                for fw_data in frameworks_data:
                    self.stdout.write(f"  - {fw_data.get('name', 'Unknown')}")
                    
        except FileNotFoundError:
            raise CommandError(f'File not found: {docx_path}')
        except Exception as e:
            raise CommandError(f'Error importing document: {str(e)}')

    def parse_document(self, doc):
        """
        Parse the Word document to extract framework data.
        This is a basic implementation - you may need to customize based on your document structure.
        """
        frameworks_data = []
        current_framework = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Try to detect framework headers (customize based on your document format)
            # Example patterns: "Chen et al. 2019", "Framework: Li et al. 2023"
            framework_match = re.search(r'([A-Z][a-z]+(?:\s+et\s+al\.)?)\s*(\d{4})?', text, re.IGNORECASE)
            if framework_match:
                if current_framework:
                    frameworks_data.append(current_framework)
                
                authors = framework_match.group(1)
                year = int(framework_match.group(2)) if framework_match.group(2) else None
                
                current_framework = {
                    'name': f"{authors} {year}" if year else authors,
                    'authors': authors,
                    'year': year,
                    'title': '',
                    'description': '',
                    'source': '',
                    'criteria': [],
                }
            elif current_framework:
                # Try to detect criteria (customize based on your document format)
                # Look for common criterion names
                criterion_patterns = [
                    r'(Completeness|Accuracy|Consistency|Conciseness|Timeliness|Relevancy|Interoperability|Availability|Usability)',
                    r'Criterion[:\s]+([A-Z][a-z]+)',
                ]
                
                for pattern in criterion_patterns:
                    match = re.search(pattern, text, re.IGNORECASE)
                    if match:
                        criterion_name = match.group(1) if match.groups() else match.group(0)
                        current_framework['criteria'].append({
                            'name': criterion_name.strip(),
                            'description': text,
                            'category': '',
                            'definitions': [text] if len(text) > 50 else [],
                        })
                        break
        
        if current_framework:
            frameworks_data.append(current_framework)
        
        # Also try to parse tables
        for table in doc.tables:
            frameworks_from_table = self.parse_table(table)
            frameworks_data.extend(frameworks_from_table)
        
        return frameworks_data

    def parse_table(self, table):
        """
        Parse tables in the document.
        Assumes tables have headers and contain framework/criterion data.
        """
        frameworks_data = []
        
        if not table.rows:
            return frameworks_data
        
        # Try to detect header row
        header_row = table.rows[0]
        headers = [cell.text.strip() for cell in header_row.cells]
        
        # Look for common column names
        framework_col = None
        criterion_col = None
        definition_col = None
        
        for i, header in enumerate(headers):
            header_lower = header.lower()
            if 'framework' in header_lower or 'author' in header_lower:
                framework_col = i
            elif 'criterion' in header_lower or 'metric' in header_lower:
                criterion_col = i
            elif 'definition' in header_lower or 'description' in header_lower:
                definition_col = i
        
        # Group rows by framework
        current_framework = None
        for row in table.rows[1:]:  # Skip header
            cells = [cell.text.strip() for cell in row.cells]
            
            if framework_col is not None and framework_col < len(cells):
                framework_name = cells[framework_col]
                if framework_name:
                    if current_framework:
                        frameworks_data.append(current_framework)
                    
                    # Extract year from framework name if present
                    year_match = re.search(r'(\d{4})', framework_name)
                    year = int(year_match.group(1)) if year_match else None
                    
                    current_framework = {
                        'name': framework_name,
                        'authors': framework_name.split()[0] if framework_name else '',
                        'year': year,
                        'title': '',
                        'description': '',
                        'source': '',
                        'criteria': [],
                    }
            
            if current_framework and criterion_col is not None and criterion_col < len(cells):
                criterion_name = cells[criterion_col] if criterion_col < len(cells) else ''
                definition_text = cells[definition_col] if definition_col is not None and definition_col < len(cells) else ''
                
                if criterion_name:
                    current_framework['criteria'].append({
                        'name': criterion_name,
                        'description': definition_text,
                        'category': '',
                        'definitions': [definition_text] if definition_text else [],
                    })
        
        if current_framework:
            frameworks_data.append(current_framework)
        
        return frameworks_data

    def import_frameworks(self, frameworks_data):
        """Import frameworks data into the database"""
        imported_count = 0
        
        for fw_data in frameworks_data:
            # Create or get framework
            framework, created = Framework.objects.get_or_create(
                name=fw_data['name'],
                defaults={
                    'authors': fw_data.get('authors', ''),
                    'year': fw_data.get('year'),
                    'title': fw_data.get('title', ''),
                    'description': fw_data.get('description', ''),
                    'source': fw_data.get('source', ''),
                }
            )
            
            if created:
                imported_count += 1
                self.stdout.write(f'Created framework: {framework.name}')
            
            # Import criteria
            for idx, criterion_data in enumerate(fw_data.get('criteria', [])):
                criterion, _ = Criterion.objects.get_or_create(
                    framework=framework,
                    name=criterion_data['name'],
                    defaults={
                        'description': criterion_data.get('description', ''),
                        'category': criterion_data.get('category', ''),
                        'order': idx,
                    }
                )
                
                # Import definitions
                for definition_text in criterion_data.get('definitions', []):
                    if definition_text.strip():
                        Definition.objects.get_or_create(
                            criterion=criterion,
                            definition_text=definition_text.strip(),
                            defaults={
                                'notes': '',
                            }
                        )
        
        return imported_count
