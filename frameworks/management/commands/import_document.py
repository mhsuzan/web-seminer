"""
Management command to import Knowledge Graph quality frameworks data from Word documents or PDF files.

Usage:
    python manage.py import_document path/to/document.docx
    python manage.py import_document path/to/document.pdf

The document should contain tables or structured data with framework information.
Supports both .docx and .pdf formats.
"""
from django.core.management.base import BaseCommand, CommandError
from django.db import transaction
import os
import re
from frameworks.models import Framework, Criterion, Definition

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False


class Command(BaseCommand):
    help = 'Import Knowledge Graph quality frameworks from a Word document (.docx) or PDF file (.pdf)'

    def add_arguments(self, parser):
        parser.add_argument('document_file', type=str, help='Path to the document (.docx or .pdf)')
        parser.add_argument(
            '--dry-run',
            action='store_true',
            help='Run without actually saving data to database',
        )

    def detect_file_type(self, file_path):
        """Detect the actual file type by reading file header"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
                # DOCX files start with PK (ZIP signature)
                if header.startswith(b'PK'):
                    return 'docx'
                # PDF files start with %PDF
                elif header.startswith(b'%PDF'):
                    return 'pdf'
        except:
            pass
        return None

    def handle(self, *args, **options):
        document_path = options['document_file']
        dry_run = options['dry_run']

        if not os.path.exists(document_path):
            raise CommandError(f'File not found: {document_path}')

        file_ext = os.path.splitext(document_path)[1].lower()
        
        # Detect actual file type (in case file extension doesn't match)
        actual_type = self.detect_file_type(document_path)
        
        try:
            if file_ext == '.docx' or actual_type == 'docx':
                if not DOCX_AVAILABLE:
                    raise CommandError('python-docx is not installed. Install it with: pip install python-docx')
                doc = Document(document_path)
                self.stdout.write(self.style.SUCCESS(f'Opened DOCX document: {document_path}'))
                frameworks_data = self.parse_docx(doc)
            elif file_ext == '.pdf' or actual_type == 'pdf':
                if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
                    raise CommandError('PDF library is not installed. Install with: pip install pdfplumber PyPDF2')
                self.stdout.write(self.style.SUCCESS(f'Opened PDF document: {document_path}'))
                frameworks_data = self.parse_pdf(document_path)
            else:
                # Try DOCX first if extension is unknown
                if DOCX_AVAILABLE:
                    try:
                        doc = Document(document_path)
                        self.stdout.write(self.style.SUCCESS(f'Detected DOCX format: {document_path}'))
                        frameworks_data = self.parse_docx(doc)
                    except:
                        raise CommandError(f'Unsupported file format: {file_ext}. Supported formats: .docx, .pdf')
                else:
                    raise CommandError(f'Unsupported file format: {file_ext}. Supported formats: .docx, .pdf')
            
            if dry_run:
                self.stdout.write(self.style.WARNING('DRY RUN MODE - No data will be saved'))
            
            self.stdout.write(self.style.SUCCESS(f'Found {len(frameworks_data)} frameworks'))
            
            if not dry_run:
                with transaction.atomic():
                    imported_count = self.import_frameworks(frameworks_data)
                    self.stdout.write(
                        self.style.SUCCESS(f'Successfully imported {imported_count} frameworks')
                    )
            else:
                self.stdout.write('Would import:')
                for fw_data in frameworks_data:
                    self.stdout.write(f"  - {fw_data.get('name', 'Unknown')}")
                    
        except Exception as e:
            raise CommandError(f'Error importing document: {str(e)}')

    def parse_docx(self, doc):
        """Parse DOCX document to extract framework data"""
        frameworks_data = []
        
        # Store document reference for hyperlink extraction
        self.doc = doc
        
        # Parse tables first (more reliable for structured data)
        # This document uses tables, so we prioritize table parsing
        for table in doc.tables:
            frameworks_from_table = self.parse_table(table)
            frameworks_data.extend(frameworks_from_table)
        
        # Only parse paragraphs if no tables found
        if not doc.tables:
            current_framework = None
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Skip document headers
                if text.lower().startswith(('comprehensive', 'the following', 'table present', 'literature review')):
                    continue
                
                # Try to detect framework headers
                framework_match = re.search(r'([A-Z][a-z]+(?:\s+et\s+al\.)?)\s*(\d{4})?', text, re.IGNORECASE)
                if framework_match:
                    if current_framework:
                        frameworks_data.append(current_framework)
                    
                    authors = framework_match.group(1)
                    year = int(framework_match.group(2)) if framework_match.group(2) else None
                    
                    current_framework = {
                        'name': f"{authors} {year}" if year else authors,
                        'authors': authors,
                        'year': year,
                        'title': '',
                        'description': '',
                        'source': '',
                        'criteria': [],
                    }
                elif current_framework:
                    # Try to detect criteria
                    criterion_patterns = [
                        r'(Completeness|Accuracy|Consistency|Conciseness|Timeliness|Relevancy|Interoperability|Availability|Usability)',
                        r'Criterion[:\s]+([A-Z][a-z]+)',
                    ]
                    
                    for pattern in criterion_patterns:
                        match = re.search(pattern, text, re.IGNORECASE)
                        if match:
                            criterion_name = match.group(1) if match.groups() else match.group(0)
                            current_framework['criteria'].append({
                                'name': criterion_name.strip(),
                                'description': text,
                                'category': '',
                                'definitions': [text] if len(text) > 50 else [],
                            })
                            break
            
            if current_framework:
                frameworks_data.append(current_framework)
        
        return frameworks_data

    def parse_pdf(self, pdf_path):
        """Parse PDF document to extract framework data"""
        frameworks_data = []
        
        # Try pdfplumber first (better for tables)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    all_text = []
                    tables_data = []
                    
                    for page in pdf.pages:
                        # Extract text
                        text = page.extract_text()
                        if text:
                            all_text.append(text)
                        
                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            tables_data.extend(tables)
                    
                    # Parse text content
                    full_text = '\n'.join(all_text)
                    frameworks_from_text = self.parse_text_content(full_text)
                    frameworks_data.extend(frameworks_from_text)
                    
                    # Parse tables
                    for table in tables_data:
                        frameworks_from_table = self.parse_pdf_table(table)
                        frameworks_data.extend(frameworks_from_table)
                    
                    return frameworks_data
            except Exception as e:
                self.stdout.write(self.style.WARNING(f'pdfplumber parsing failed: {e}, trying PyPDF2'))
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    all_text = []
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            text = page.extract_text()
                            if text:
                                all_text.append(text)
                        except Exception as e:
                            self.stdout.write(self.style.WARNING(f'Error extracting text from page {page_num + 1}: {e}'))
                            continue
                    
                    full_text = '\n'.join(all_text)
                    if not full_text.strip():
                        raise CommandError('No text could be extracted from PDF. The PDF might be image-based or corrupted.')
                    frameworks_data = self.parse_text_content(full_text)
                    return frameworks_data
            except Exception as e:
                raise CommandError(f'Failed to parse PDF with PyPDF2: {e}')
        
        raise CommandError('No PDF parsing library available')

    def parse_text_content(self, text):
        """Parse text content to extract framework data"""
        frameworks_data = []
        current_framework = None
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Try to detect framework headers
            # Pattern: "Author et al. (Year)" or "Author (Year)" or "Framework Name"
            framework_patterns = [
                r'([A-Z][a-z]+(?:\s+et\s+al\.)?)\s*\(?(\d{4})\)?',
                r'Framework[:\s]+([A-Z][^\(]+)',
                r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s+(\d{4})',
            ]
            
            for pattern in framework_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    if current_framework:
                        frameworks_data.append(current_framework)
                    
                    authors = match.group(1).strip() if match.groups() else match.group(0)
                    year = int(match.group(2)) if len(match.groups()) > 1 and match.group(2) else None
                    
                    current_framework = {
                        'name': f"{authors} {year}" if year else authors,
                        'authors': authors,
                        'year': year,
                        'title': '',
                        'description': '',
                        'source': '',
                        'criteria': [],
                    }
                    break
            
            # Try to detect criteria
            if current_framework:
                criterion_patterns = [
                    r'^\s*[-•]\s*(Completeness|Accuracy|Consistency|Conciseness|Timeliness|Relevancy|Interoperability|Availability|Usability|Correctness|Currency|Coverage)',
                    r'(Completeness|Accuracy|Consistency|Conciseness|Timeliness|Relevancy|Interoperability|Availability|Usability|Correctness|Currency|Coverage)[:\s]+',
                    r'Criterion[:\s]+([A-Z][a-z]+)',
                    r'^\s*\d+\.\s*([A-Z][a-z]+)',
                ]
                
                for pattern in criterion_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        criterion_name = match.group(1) if match.groups() else match.group(0)
                        # Get description (rest of the line or next lines)
                        description = line.replace(match.group(0), '').strip()
                        
                        current_framework['criteria'].append({
                            'name': criterion_name.strip(),
                            'description': description,
                            'category': '',
                            'definitions': [description] if description else [],
                        })
                        break
        
        if current_framework:
            frameworks_data.append(current_framework)
        
        return frameworks_data

    def parse_pdf_table(self, table):
        """Parse PDF table to extract framework data"""
        frameworks_data = []
        
        if not table or len(table) == 0:
            return frameworks_data
        
        # Try to detect header row
        header_row = table[0] if len(table) > 0 else []
        headers = [str(cell).strip() if cell else '' for cell in header_row]
        
        # Look for common column names
        framework_col = None
        criterion_col = None
        definition_col = None
        
        for i, header in enumerate(headers):
            header_lower = str(header).lower()
            if 'framework' in header_lower or 'author' in header_lower or 'source' in header_lower:
                framework_col = i
            elif 'criterion' in header_lower or 'metric' in header_lower or 'dimension' in header_lower:
                criterion_col = i
            elif 'definition' in header_lower or 'description' in header_lower:
                definition_col = i
        
        # Group rows by framework
        current_framework = None
        for row in table[1:]:  # Skip header
            if not row:
                continue
            
            cells = [str(cell).strip() if cell else '' for cell in row]
            
            if framework_col is not None and framework_col < len(cells):
                framework_name = cells[framework_col]
                if framework_name:
                    if current_framework:
                        frameworks_data.append(current_framework)
                    
                    # Extract year from framework name if present
                    year_match = re.search(r'(\d{4})', framework_name)
                    year = int(year_match.group(1)) if year_match else None
                    
                    current_framework = {
                        'name': framework_name,
                        'authors': framework_name.split()[0] if framework_name else '',
                        'year': year,
                        'title': '',
                        'description': '',
                        'source': '',
                        'criteria': [],
                    }
            
            if current_framework and criterion_col is not None and criterion_col < len(cells):
                criterion_name = cells[criterion_col] if criterion_col < len(cells) else ''
                definition_text = cells[definition_col] if definition_col is not None and definition_col < len(cells) else ''
                
                if criterion_name:
                    current_framework['criteria'].append({
                        'name': criterion_name,
                        'description': definition_text,
                        'category': '',
                        'definitions': [definition_text] if definition_text else [],
                    })
        
        if current_framework:
            frameworks_data.append(current_framework)
        
        return frameworks_data

    def extract_hyperlinks_from_cell(self, cell):
        """
        Extract hyperlinks from a DOCX cell.
        Returns a tuple: (text, url) where url is the first hyperlink found, or None if no hyperlink.
        """
        text = cell.text.strip()
        url = None
        
        try:
            # Get the document part for accessing relationships
            doc_part = None
            if hasattr(self, 'doc') and hasattr(self.doc, 'part'):
                doc_part = self.doc.part
            
            # Iterate through paragraphs - hyperlinks are often at paragraph level
            for paragraph in cell.paragraphs:
                part = paragraph.part if hasattr(paragraph, 'part') else doc_part
                if not part:
                    continue
                
                # Check paragraph element for hyperlink children
                para_elem = paragraph._element
                for child in para_elem:
                    # Check if this child is a hyperlink element
                    if child.tag and 'hyperlink' in child.tag:
                        # Get relationship ID (rId) - this is for external links
                        r_id = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if r_id and part and hasattr(part, 'rels'):
                            try:
                                rel = part.rels.get(r_id)
                                if rel:
                                    # Get the target URL - try different attribute names
                                    if hasattr(rel, 'target_ref'):
                                        url = rel.target_ref
                                    elif hasattr(rel, '_target'):
                                        url = str(rel._target)
                                    elif hasattr(rel, 'target_uri'):
                                        url = str(rel.target_uri)
                                    elif hasattr(rel, 'target'):
                                        url = str(rel.target)
                                    
                                    # If we found a URL, return it immediately
                                    if url:
                                        return (text, url)
                            except Exception as e:
                                # Continue trying other methods
                                pass
                
                # Also check runs for hyperlinks (some documents might have them nested differently)
                for run in paragraph.runs:
                    run_elem = run._element
                    
                    # Look for hyperlink elements in the run (recursively)
                    def find_hyperlinks(elem):
                        """Recursively find hyperlink elements"""
                        hyperlinks = []
                        # Check if this element is a hyperlink
                        if elem.tag and 'hyperlink' in elem.tag:
                            hyperlinks.append(elem)
                        # Check children
                        for child in elem:
                            hyperlinks.extend(find_hyperlinks(child))
                        return hyperlinks
                    
                    hyperlinks = find_hyperlinks(run_elem)
                    
                    for hyperlink in hyperlinks:
                        # Check for relationship ID (rId) - this is for external links
                        r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if r_id and part and hasattr(part, 'rels'):
                            try:
                                rel = part.rels.get(r_id)
                                if rel:
                                    # Get the target URL - try different attribute names
                                    if hasattr(rel, 'target_ref'):
                                        url = rel.target_ref
                                    elif hasattr(rel, '_target'):
                                        url = str(rel._target)
                                    elif hasattr(rel, 'target_uri'):
                                        url = str(rel.target_uri)
                                    elif hasattr(rel, 'target'):
                                        url = str(rel.target)
                                    
                                    # If we found a URL, return it immediately
                                    if url:
                                        return (text, url)
                            except Exception as e:
                                # Continue trying other methods
                                pass
                        
        except Exception as e:
            # If hyperlink extraction fails, just return the text
            # This is a fallback to ensure the import doesn't break
            pass
        
        return (text, url)
    
    def parse_table(self, table):
        """Parse DOCX table to extract framework data"""
        frameworks_data = []
        
        if not table.rows or len(table.rows) < 2:
            return frameworks_data
        
        # Get header row
        header_row = table.rows[0]
        headers = [cell.text.strip().lower() for cell in header_row.cells]
        
        # Find column indices for our specific table structure
        title_col = None
        year_col = None
        dimensions_col = None
        abstract_col = None
        objectives_col = None
        methodology_col = None
        algorithm_col = None
        top_model_col = None
        accuracy_col = None
        advantages_col = None
        drawbacks_col = None
        reference_col = None
        
        for i, header in enumerate(headers):
            if 'title' in header:
                title_col = i
            elif 'year' in header or 'published' in header:
                year_col = i
            elif 'dimension' in header:
                dimensions_col = i
            elif 'abstract' in header:
                abstract_col = i
            elif 'objective' in header:
                objectives_col = i
            elif 'methodology' in header:
                methodology_col = i
            elif 'algorithm' in header:
                algorithm_col = i
            elif 'top model' in header or 'topmodel' in header.replace(' ', '').lower():
                top_model_col = i
            elif 'accuracy' in header:
                accuracy_col = i
            elif 'advantage' in header:
                advantages_col = i
            elif 'drawback' in header:
                drawbacks_col = i
            elif 'reference' in header:
                reference_col = i
        
        # Parse each data row
        for row in table.rows[1:]:  # Skip header
            # Extract text from all cells first
            cells = [cell.text.strip() for cell in row.cells]
            
            # Extract framework information
            title = cells[title_col] if title_col is not None and title_col < len(cells) else ''
            year_str = cells[year_col] if year_col is not None and year_col < len(cells) else ''
            dimensions = cells[dimensions_col] if dimensions_col is not None and dimensions_col < len(cells) else ''
            abstract = cells[abstract_col] if abstract_col is not None and abstract_col < len(cells) else ''
            objectives = cells[objectives_col] if objectives_col is not None and objectives_col < len(cells) else ''
            methodology = cells[methodology_col] if methodology_col is not None and methodology_col < len(cells) else ''
            algorithm_used = cells[algorithm_col] if algorithm_col is not None and algorithm_col < len(cells) else ''
            top_model = cells[top_model_col] if top_model_col is not None and top_model_col < len(cells) else ''
            accuracy = cells[accuracy_col] if accuracy_col is not None and accuracy_col < len(cells) else ''
            advantages = cells[advantages_col] if advantages_col is not None and advantages_col < len(cells) else ''
            drawbacks = cells[drawbacks_col] if drawbacks_col is not None and drawbacks_col < len(cells) else ''
            
            # Extract reference with hyperlink support
            reference = ''
            if reference_col is not None and reference_col < len(row.cells):
                ref_cell = row.cells[reference_col]
                ref_text, ref_url = self.extract_hyperlinks_from_cell(ref_cell)
                
                # Combine text and URL appropriately
                if ref_url:
                    # If we have a URL, use it (prefer URL over text if text is just "Read" or similar)
                    if ref_text.lower() in ['read', 'link', 'url', 'source']:
                        reference = ref_url
                    else:
                        # If text is meaningful, combine them: "text (URL)" or just URL if text is empty
                        if ref_text and ref_text.strip():
                            combined = f"{ref_text} ({ref_url})"
                            # Ensure we don't exceed the 500 character limit for the source field
                            if len(combined) > 500:
                                # If combined is too long, prefer the URL
                                reference = ref_url[:500]
                            else:
                                reference = combined
                        else:
                            reference = ref_url[:500]  # Truncate if needed
                else:
                    # No hyperlink found, just use the text
                    reference = ref_text[:500] if ref_text else ''  # Truncate if needed
            
            # Skip if title is too short or looks like a document header
            if not title or len(title) < 5:
                continue
            
            # Skip document headers
            title_lower = title.lower()
            if any(title_lower.startswith(prefix) for prefix in ('comprehensive', 'the following', 'table present', 'table ', 'figure ', 'page ')):
                continue
            
            # Extract year
            year = None
            if year_str:
                year_match = re.search(r'(\d{4})', year_str)
                if year_match:
                    try:
                        year = int(year_match.group(1))
                    except ValueError:
                        pass
            
            # If no year found in year column, try to extract from title
            if not year:
                year_match = re.search(r'\((\d{4})\)', title)
                if year_match:
                    try:
                        year = int(year_match.group(1))
                    except ValueError:
                        pass
            
            # Extract authors from title or reference
            # Since reference column just says "Read", we'll try to extract from title
            # or leave empty if title doesn't contain author info
            authors = ''
            
            # Try reference column first (though it usually just says "Read")
            if reference and reference.lower() != 'read':
                author_match = re.search(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*(?:\s+et\s+al\.)?)', reference)
                if author_match:
                    authors = author_match.group(1).strip()
            
            # If no authors from reference, try title patterns
            if not authors and title:
                # Look for common author patterns in titles
                # Pattern: "Author et al. - Title" or "Author: Title"
                title_clean = re.sub(r'\s*\(?\d{4}\)?', '', title)
                
                # Check if title starts with what looks like an author name (short, capitalized words)
                parts = re.split(r'[:\-–]', title_clean, 1)
                if parts and len(parts) > 0:
                    first_part = parts[0].strip()
                    words = first_part.split()
                    # If first part is short (likely author), use it
                    if len(words) <= 4 and len(first_part) < 50:
                        # Check if it looks like an author name (starts with capital, has 2-4 words)
                        if all(w[0].isupper() if w else False for w in words[:2]):
                            authors = first_part
                
                # If still no authors, leave empty (will be stored as empty string)
            
            # Parse dimensions/criteria
            criteria = []
            if dimensions:
                # Normalize the dimensions string - replace newlines with spaces first
                dimensions_normalized = re.sub(r'\s+', ' ', dimensions)
                
                # Handle special cases where words are split (e.g., "Syntactic\nValidity" -> "Syntactic Validity")
                # Join words that might have been split: "Syntactic Validity", "Semantic Accuracy", etc.
                dimensions_normalized = re.sub(r'\b(Syntactic|Semantic|Representational)\s+([A-Z][a-z]+)', r'\1 \2', dimensions_normalized)
                
                # Split by comma, semicolon, or newline
                dim_list = re.split(r'[,;\n]+', dimensions_normalized)
                
                seen_dimensions = set()  # Avoid duplicates
                
                for dim in dim_list:
                    dim = dim.strip()
                    # Filter out very short strings and common non-dimension words
                    if dim and len(dim) > 2 and dim.lower() not in ['n/a', 'na', 'read', 'and', 'or', 'the', 'none', 'null']:
                        # Clean up common prefixes that might be split across lines
                        dim = re.sub(r'^(Syntactic|Semantic|Representational)[\s-]+', '', dim, flags=re.IGNORECASE)
                        # Remove trailing periods, dashes, and parentheses
                        dim = dim.rstrip('.-()[]').strip()
                        
                        # Skip if it's just a single letter, number, or common words
                        if dim and len(dim) > 2 and not re.match(r'^[\d\s]+$', dim):
                            # Capitalize first letter
                            dim = dim[0].upper() + dim[1:] if len(dim) > 1 else dim
                            
                            # Avoid duplicates (case-insensitive)
                            dim_lower = dim.lower()
                            if dim_lower not in seen_dimensions:
                                seen_dimensions.add(dim_lower)
                                # Use abstract or description if available for better definition
                                definition_text = abstract if abstract else f"Quality dimension from {title}"
                                if year:
                                    definition_text += f" ({year})"
                                
                                criteria.append({
                                    'name': dim,
                                    'description': definition_text,
                                    'category': '',
                                    'definitions': [definition_text],
                                })
            
            # Create framework entry
            framework_data = {
                'name': title,
                'authors': authors,
                'year': year,
                'title': title,
                'description': abstract if abstract else '',
                'objectives': objectives if objectives else '',
                'methodology': methodology if methodology else '',
                'algorithm_used': algorithm_used if algorithm_used else '',
                'top_model': top_model if top_model else '',
                'accuracy': accuracy if accuracy else '',
                'advantages': advantages if advantages else '',
                'drawbacks': drawbacks if drawbacks else '',
                'source': reference if reference else '',
                'criteria': criteria,
            }
            
            frameworks_data.append(framework_data)
        
        return frameworks_data

    def normalize_name(self, name):
        """Normalize a name for comparison (lowercase, strip, remove extra spaces)"""
        if not name:
            return ''
        return ' '.join(name.lower().strip().split())
    
    def find_matching_framework(self, fw_data):
        """Find existing framework by normalized name, year, or title"""
        name = fw_data.get('name', '').strip()
        year = fw_data.get('year')
        title = fw_data.get('title', '').strip()
        
        # Try exact name match first
        framework = Framework.objects.filter(name=name).first()
        if framework:
            return framework
        
        # Try normalized name match
        normalized_name = self.normalize_name(name)
        if normalized_name:
            for fw in Framework.objects.all():
                if self.normalize_name(fw.name) == normalized_name:
                    return fw
        
        # Try matching by year and title (if both exist)
        if year and title:
            framework = Framework.objects.filter(year=year, title=title).first()
            if framework:
                return framework
        
        # Try matching by year and normalized title
        if year and title:
            normalized_title = self.normalize_name(title)
            for fw in Framework.objects.filter(year=year):
                if self.normalize_name(fw.title) == normalized_title:
                    return fw
        
        return None
    
    def merge_framework_data(self, framework, fw_data):
        """Merge new data into existing framework, keeping existing data if new is empty"""
        updated = False
        
        # Only update if new data is not empty and different
        if fw_data.get('authors') and fw_data['authors'].strip() and (not framework.authors or framework.authors.strip() != fw_data['authors'].strip()):
            framework.authors = fw_data['authors'].strip()
            updated = True
        
        if fw_data.get('year') and framework.year != fw_data['year']:
            framework.year = fw_data['year']
            updated = True
        
        if fw_data.get('title') and fw_data['title'].strip() and (not framework.title or framework.title.strip() != fw_data['title'].strip()):
            framework.title = fw_data['title'].strip()
            updated = True
        
        if fw_data.get('description') and fw_data['description'].strip() and (not framework.description or len(fw_data['description'].strip()) > len(framework.description.strip())):
            framework.description = fw_data['description'].strip()
            updated = True
        
        if fw_data.get('objectives') and fw_data['objectives'].strip() and (not framework.objectives or len(fw_data['objectives'].strip()) > len(framework.objectives.strip())):
            framework.objectives = fw_data['objectives'].strip()
            updated = True
        
        if fw_data.get('methodology') and fw_data['methodology'].strip() and (not framework.methodology or len(fw_data['methodology'].strip()) > len(framework.methodology.strip())):
            framework.methodology = fw_data['methodology'].strip()
            updated = True
        
        if fw_data.get('algorithm_used') and fw_data['algorithm_used'].strip() and (not framework.algorithm_used or framework.algorithm_used.strip() != fw_data['algorithm_used'].strip()):
            framework.algorithm_used = fw_data['algorithm_used'].strip()
            updated = True
        
        if fw_data.get('top_model') and fw_data['top_model'].strip() and (not framework.top_model or framework.top_model.strip() != fw_data['top_model'].strip()):
            framework.top_model = fw_data['top_model'].strip()
            updated = True
        
        if fw_data.get('accuracy') and fw_data['accuracy'].strip() and (not framework.accuracy or framework.accuracy.strip() != fw_data['accuracy'].strip()):
            framework.accuracy = fw_data['accuracy'].strip()
            updated = True
        
        if fw_data.get('advantages') and fw_data['advantages'].strip() and (not framework.advantages or len(fw_data['advantages'].strip()) > len(framework.advantages.strip())):
            framework.advantages = fw_data['advantages'].strip()
            updated = True
        
        if fw_data.get('drawbacks') and fw_data['drawbacks'].strip() and (not framework.drawbacks or len(fw_data['drawbacks'].strip()) > len(framework.drawbacks.strip())):
            framework.drawbacks = fw_data['drawbacks'].strip()
            updated = True
        
        if fw_data.get('source') and fw_data['source'].strip():
            new_source = fw_data['source'].strip()
            current_source = framework.source.strip() if framework.source else ''
            
            # Prefer URLs over plain text like "Read"
            # If new source looks like a URL and current doesn't, update
            is_url = new_source.startswith(('http://', 'https://', 'www.'))
            current_is_url = current_source.startswith(('http://', 'https://', 'www.'))
            
            # Update if:
            # 1. Current source is empty, OR
            # 2. New source is a URL and current is not, OR
            # 3. New source is different (and both are same type or new is better)
            should_update = (not current_source or 
                           (is_url and not current_is_url) or
                           (new_source != current_source and not (current_is_url and not is_url)))
            
            if should_update:
                framework.source = new_source
                updated = True
        
        if updated:
            framework.save()
        
        return updated
    
    def normalize_criterion_name(self, name):
        """Normalize criterion name for comparison"""
        if not name:
            return ''
        # Remove extra whitespace, normalize case
        normalized = ' '.join(name.strip().split())
        # Capitalize first letter
        if normalized:
            normalized = normalized[0].upper() + normalized[1:] if len(normalized) > 1 else normalized.upper()
        return normalized
    
    def find_matching_criterion(self, framework, criterion_name):
        """Find existing criterion by normalized name"""
        normalized_name = self.normalize_criterion_name(criterion_name)
        if not normalized_name:
            return None
        
        # Try exact match first
        criterion = Criterion.objects.filter(framework=framework, name=criterion_name).first()
        if criterion:
            return criterion
        
        # Try normalized match
        for crit in Criterion.objects.filter(framework=framework):
            if self.normalize_criterion_name(crit.name) == normalized_name:
                return crit
        
        return None
    
    def import_frameworks(self, frameworks_data):
        """Import frameworks data into the database with duplicate detection"""
        imported_count = 0
        updated_count = 0
        
        for fw_data in frameworks_data:
            # Normalize framework name
            fw_data['name'] = fw_data.get('name', '').strip()
            if not fw_data['name']:
                self.stdout.write(self.style.WARNING('Skipping framework with empty name'))
                continue
            
            # Try to find existing framework
            framework = self.find_matching_framework(fw_data)
            
            if framework:
                # Update existing framework
                was_updated = self.merge_framework_data(framework, fw_data)
                if was_updated:
                    updated_count += 1
                    self.stdout.write(f'Updated framework: {framework.name}')
            else:
                # Create new framework
                framework = Framework.objects.create(
                    name=fw_data['name'],
                    authors=fw_data.get('authors', '').strip(),
                    year=fw_data.get('year'),
                    title=fw_data.get('title', '').strip(),
                    description=fw_data.get('description', '').strip(),
                    objectives=fw_data.get('objectives', '').strip(),
                    methodology=fw_data.get('methodology', '').strip(),
                    algorithm_used=fw_data.get('algorithm_used', '').strip(),
                    top_model=fw_data.get('top_model', '').strip(),
                    accuracy=fw_data.get('accuracy', '').strip(),
                    advantages=fw_data.get('advantages', '').strip(),
                    drawbacks=fw_data.get('drawbacks', '').strip(),
                    source=fw_data.get('source', '').strip(),
                )
                imported_count += 1
                self.stdout.write(f'Created framework: {framework.name}')
            
            # Import criteria with duplicate detection
            for idx, criterion_data in enumerate(fw_data.get('criteria', [])):
                criterion_name = criterion_data.get('name', '').strip()
                if not criterion_name:
                    continue
                
                # Normalize criterion name
                normalized_name = self.normalize_criterion_name(criterion_name)
                
                # Try to find existing criterion
                criterion = self.find_matching_criterion(framework, criterion_name)
                
                if criterion:
                    # Update existing criterion if new data is better
                    if criterion_data.get('description') and criterion_data['description'].strip():
                        if not criterion.description or len(criterion_data['description'].strip()) > len(criterion.description.strip()):
                            criterion.description = criterion_data['description'].strip()
                            criterion.save()
                    if criterion_data.get('category') and criterion_data['category'].strip():
                        if not criterion.category or criterion.category.strip() != criterion_data['category'].strip():
                            criterion.category = criterion_data['category'].strip()
                            criterion.save()
                else:
                    # Create new criterion
                    criterion = Criterion.objects.create(
                        framework=framework,
                        name=normalized_name,
                        description=criterion_data.get('description', '').strip(),
                        category=criterion_data.get('category', '').strip(),
                        order=idx,
                    )
                
                # Import definitions with duplicate detection
                for definition_text in criterion_data.get('definitions', []):
                    definition_text = definition_text.strip()
                    if not definition_text:
                        continue
                    
                    # Check if similar definition already exists (normalized comparison)
                    normalized_def = self.normalize_name(definition_text)
                    existing_def = None
                    
                    for def_obj in criterion.definitions.all():
                        if self.normalize_name(def_obj.definition_text) == normalized_def:
                            existing_def = def_obj
                            break
                    
                    if not existing_def:
                        # Only create if significantly different (avoid near-duplicates)
                        is_duplicate = False
                        for def_obj in criterion.definitions.all():
                            existing_normalized = self.normalize_name(def_obj.definition_text)
                            # Check if one is a substring of the other (likely duplicate)
                            if normalized_def in existing_normalized or existing_normalized in normalized_def:
                                if abs(len(normalized_def) - len(existing_normalized)) < 20:  # Similar length
                                    is_duplicate = True
                                    break
                        
                        if not is_duplicate:
                            Definition.objects.create(
                                criterion=criterion,
                                definition_text=definition_text,
                                notes='',
                            )
        
        self.stdout.write(self.style.SUCCESS(f'Imported {imported_count} new frameworks, updated {updated_count} existing frameworks'))
        return imported_count
